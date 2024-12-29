# Libraries for markdown formatting
import textwrap
from IPython.display import display, Markdown

# Libraries for word document generation
import markdown2
from docx import Document

import re

def print_formatted_md(md_text: str, max_line_height: int = 80):
    """
    Formats and prints Markdown text with each line wrapped to the specified width.

    Args:
    md_text (str): The Markdown text to be formatted and printed.
    max_line_height (int): The maximum number of characters per line. Default is 80.
    """
    lines = md_text.splitlines()
    wrapped_lines = []

    for line in lines:
        # Handle list items and pragraphs with bold text
        if line.startswith(('* ', '- ', '> ', '**')):
            prefix = ""
            if line.startswith('**'):
                prefix = "**"
                line = line[2:] # Remove the first two characters (**)
            elif line.startswith('* '):
                prefix = "* "
                line = line[2:]
            elif line.startswith('- '):
                prefix = "- "
                line = line[2:]
            elif line.startswith('> '):
                prefix = "> "
                line = line[2:]
            wrapped = textwrap.wrap(line, width=max_line_height - len(prefix))
            for i, wrap in enumerate(wrapped):
                wrapped_lines.append(f"{prefix}{wrap}" if i == 0 else f"{' ' * len(prefix)}{wrap}") 
        else:
            # For normal lines, just wrap without any prefix
            wrapped_lines.extend(textwrap.wrap(line, width=max_line_height))

    # Join the wrapped lines back into a single string
    wrapped_text = "\n".join(wrapped_lines)

    # Display the formatted text for terminal
    print(wrapped_text)

    # Display the formatted text for Juptyer Notebook
    # display(Markdown(wrapped_text))

def strip_html_tags(text: str) -> str:
    """Remove HTML tags from the text."""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)

def handle_strong_tags(text: str) -> list[str]:
    """Handle <strong> tags in text and return formatted runs."""
    segments = text.split('<strong>')
    runs = []
    for segment in segments:
        if '</strong>' in segment:
            bold_text, remain = segment.split('</strong>')
            runs.append((bold_text, True)) # Text inside <strong> tags is bold
            runs.append((remain, False)) # Text after </strong> is normal
        else:
            runs.append((segment, False)) # No bold text
    return runs

def add_markdown_to_docx(md_text: str, doc):
    """
    Converts Markdown text to a Word document format and adds it to a given document object.

    Args:
        md_text (str): The input Markdown text to be converted.
        doc: A `docx.Document` object where the converted content will be added.

    This function processes Markdown headers (<h1>, <h2>, <h3>), bold text (<strong>),
    unordered lists (<ul>, <li>), and plain text. It uses the `markdown2` library to convert
    Markdown to HTML and then parses the resulting HTML to add content to the Word document.
    """
    html = markdown2.markdown(md_text)
    lines = html.splitlines()

    for line in lines:
        line = line.strip()

        if line.startswith('<h1>') or line.startswith('<h2>') or line.startswith('<h3>'):
            # Add heading
            heading = strip_html_tags(line).strip()
            doc.add_heading(heading, level=1 if line.startswith('<h1>') else 2 if line.startswith('<h2>') else 3)
        elif "<strong>" in line:
            # Handle <strong> tags
            paragraph = doc.add_paragraph()
            for text, is_bold in handle_strong_tags(line):
                run = paragraph.add_run(strip_html_tags(text))
                run.bold = is_bold
        elif line.startswith("<ul>") or line.startswith("</ul>"):
            continue
        elif line.startswith("<li>"):
            # Add bullet list item
            paragraph = doc.add_paragraph(style='List Bullet')
            for text, is_bold in handle_strong_tags(line):
                run = paragraph.add_run(strip_html_tags(text))
                run.bold = is_bold
        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run(strip_html_tags(line))

def convert_md_to_docx(md_text: str, output_filename: str):
    """
    Converts Markdown text to a Word document file and saves it.

    Args:
        md_text (str): The input Markdown text to be converted.
        output_filename (str): The name of the output Word document file.

    This function creates a new Word document, processes the Markdown text,
    and saves the content in the specified output file.
    """
    doc = Document()
    add_markdown_to_docx(md_text, doc)
    doc.save(output_filename)